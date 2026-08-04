# -*- coding: utf-8 -*-
"""
Generator for InTechSol × Alait (COMRAD) joint presentation event.
Produces 4 Word documents:
  - Program (GE / RU)
  - Information booklet (GE / RU)
Event: Friday, 25 September 2026 — SPInTechSol + COMRAD joint presentation.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.dirname(os.path.abspath(__file__))

# Brand colours (from the InTechSol logo)
NAVY  = RGBColor(0x14, 0x24, 0x53)   # deep navy
BLUE  = RGBColor(0x1E, 0x88, 0xE5)   # bright blue
GREY  = RGBColor(0x55, 0x5B, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
NAVY_HEX  = "142453"
BLUE_HEX  = "1E88E5"
LIGHT_HEX = "EAF2FB"   # light blue row shading
BAND_HEX  = "142453"

# ----------------------------------------------------------------------------- helpers

def set_base_font(doc, font_name):
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(a), font_name)

def _set_run_font(run, name):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(a), name)

def par_shade(par, fill_hex):
    pPr = par._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)

def cell_shade(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right),
                     ("left", left), ("right", right)):
        e = OxmlElement("w:" + tag)
        e.set(qn("w:w"), str(val)); e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)

def no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)

def add_run(par, text, font, size=10.5, bold=False, italic=False, color=None):
    r = par.add_run(text)
    _set_run_font(r, font)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return r

def spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    for r in ():
        pass
    p.add_run("").font.size = Pt(pts)
    return p

def band(doc, text, font, fill=BAND_HEX, tcolor=WHITE, size=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    par_shade(p, fill)
    add_run(p, "  " + text, font, size=size, bold=True, color=tcolor)
    return p

def h2(doc, text, font, size=13, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text, font, size=size, bold=True, color=color)
    # thin rule
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), BLUE_HEX)
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p

def para(doc, text, font, size=10.5, bold=False, color=None, align=None,
         space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align is not None:
        p.alignment = align
    add_run(p, text, font, size=size, bold=bold, color=color)
    return p

def bullet(doc, text, font, size=10.5, bold_lead=None):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(3)
    add_run(p, "▸  ", font, size=size, bold=True, color=BLUE)
    if bold_lead:
        add_run(p, bold_lead, font, size=size, bold=True, color=NAVY)
        add_run(p, text, font, size=size)
    else:
        add_run(p, text, font, size=size)
    return p

def page_break(doc):
    doc.add_page_break()

def set_margins(doc, top=1.9, bottom=1.9, left=2.0, right=2.0):
    for s in doc.sections:
        s.top_margin = Cm(top); s.bottom_margin = Cm(bottom)
        s.left_margin = Cm(left); s.right_margin = Cm(right)

# ----------------------------------------------------------------------------- cover

def cover(doc, font, T, doc_type):
    # top brand line
    p = para(doc, T["org_line"], font, size=11, bold=True, color=BLUE,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, T["partner_line"], font, size=10, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=26)

    para(doc, T["event_kicker"], font, size=12, bold=True, color=BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    # Big title
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    add_run(p, T["event_title"], font, size=26, bold=True, color=NAVY)
    para(doc, T["event_subtitle"], font, size=14, bold=True, color=BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=22)

    # info box (single-cell table with shading)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell_shade(cell, LIGHT_HEX)
    set_cell_margins(cell, top=180, bottom=180, left=260, right=260)
    cell.width = Cm(14)
    lines = [
        (T["l_date"],  T["v_date"]),
        (T["l_time"],  T["v_time"]),
        (T["l_venue"], T["v_venue"]),
        (T["l_format"],T["v_format"]),
    ]
    first = True
    for lab, val in lines:
        if first:
            cp = cell.paragraphs[0]; first = False
        else:
            cp = cell.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(4)
        add_run(cp, lab + "  ", font, size=11, bold=True, color=NAVY)
        add_run(cp, val, font, size=11, color=RGBColor(0x22,0x22,0x22))
    no_table_borders(tbl)

    # doc type footer on cover
    para(doc, "", font, space_after=18)
    para(doc, doc_type, font, size=12, bold=True, color=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, T["cover_foot"], font, size=9, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER)

# ----------------------------------------------------------------------------- program

def build_program(font, T, filename):
    doc = Document()
    set_base_font(doc, font)
    set_margins(doc)
    cover(doc, font, T, T["doc_type"])
    page_break(doc)

    para(doc, T["prog_title"], font, size=18, bold=True, color=NAVY, space_after=2)
    para(doc, T["prog_sub"], font, size=10.5, color=GREY, space_after=10)

    # schedule table
    tbl = doc.add_table(rows=0, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    def add_time_row(time, text, bold=False):
        row = tbl.add_row()
        c0, c1 = row.cells
        c0.width = Cm(3.1); c1.width = Cm(12.4)
        set_cell_margins(c0); set_cell_margins(c1)
        p0 = c0.paragraphs[0]; p0.paragraph_format.space_after = Pt(2)
        add_run(p0, time, font, size=10.5, bold=True, color=BLUE)
        p1 = c1.paragraphs[0]; p1.paragraph_format.space_after = Pt(2)
        add_run(p1, text, font, size=10.5, bold=bold, color=NAVY if bold else RGBColor(0x22,0x22,0x22))

    def add_section_row(text):
        row = tbl.add_row()
        a, b = row.cells
        a.merge(b)
        cell = row.cells[0]
        cell_shade(cell, NAVY_HEX)
        set_cell_margins(cell, top=80, bottom=80)
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        add_run(p, text, font, size=11, bold=True, color=WHITE)

    for item in T["schedule"]:
        if item[0] == "SEC":
            add_section_row(item[1])
        else:
            add_time_row(item[0], item[1], bold=item[2] if len(item) > 2 else False)
    no_table_borders(tbl)

    # B2B detail
    band(doc, T["b2b_head"], font)
    para(doc, T["b2b_intro"], font, space_after=6)
    for g in T["b2b_groups"]:
        bullet(doc, g[1], font, bold_lead=g[0] + " — ")

    # notes
    band(doc, T["notes_head"], font)
    for n in T["prog_notes"]:
        bullet(doc, n, font)

    para(doc, "", font, space_after=10)
    para(doc, T["contact_short"], font, size=9.5, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER)

    path = os.path.join(OUT, filename)
    doc.save(path)
    return path

# ----------------------------------------------------------------------------- booklet

def build_booklet(font, T, filename):
    doc = Document()
    set_base_font(doc, font)
    set_margins(doc)
    cover(doc, font, T, T["booklet_doc_type"])
    page_break(doc)

    # 1. Welcome / about the event
    band(doc, T["s_event"], font)
    for p in T["about_event"]:
        para(doc, p, font)

    # 2. InTechSol / SPInTechSol
    band(doc, T["s_intechsol"], font)
    h2(doc, T["is_about_h"], font)
    para(doc, T["is_about"], font)
    h2(doc, T["is_problem_h"], font)
    for b in T["is_problem"]:
        bullet(doc, b[1], font, bold_lead=b[0] + " — ")
    h2(doc, T["is_solution_h"], font)
    para(doc, T["is_solution"], font)
    # modules two lists
    h2(doc, T["is_mod8_h"], font)
    for b in T["is_mod8"]:
        bullet(doc, b, font)
    h2(doc, T["is_mod8b_h"], font)
    for b in T["is_mod8b"]:
        bullet(doc, b, font)
    h2(doc, T["is_ai_h"], font)
    para(doc, T["is_ai"], font)
    h2(doc, T["is_std_h"], font)
    for b in T["is_std"]:
        bullet(doc, b[1], font, bold_lead=b[0] + " — ")
    h2(doc, T["is_price_h"], font)
    para(doc, T["is_price"], font)

    page_break(doc)

    # 3. Alait / COMRAD
    band(doc, T["s_comrad"], font)
    h2(doc, T["cr_about_h"], font)
    para(doc, T["cr_about"], font)
    h2(doc, T["cr_what_h"], font)
    para(doc, T["cr_what"], font)
    h2(doc, T["cr_line_h"], font)
    for b in T["cr_line"]:
        bullet(doc, b[1], font, bold_lead=b[0] + " — ")
    h2(doc, T["cr_func_h"], font)
    for b in T["cr_func"]:
        bullet(doc, b[1], font, bold_lead=b[0] + " — ")
    h2(doc, T["cr_tactic_h"], font)
    para(doc, T["cr_tactic"], font)
    h2(doc, T["cr_app_h"], font)
    for b in T["cr_app"]:
        bullet(doc, b, font)
    h2(doc, T["cr_adv_h"], font)
    for b in T["cr_adv"]:
        bullet(doc, b, font)

    page_break(doc)

    # 4. Synergy
    band(doc, T["s_synergy"], font)
    para(doc, T["synergy_intro"], font)
    for b in T["synergy"]:
        bullet(doc, b, font)

    # 5. Audience
    band(doc, T["s_audience"], font)
    para(doc, T["audience_intro"], font)
    for b in T["audience"]:
        bullet(doc, b[1], font, bold_lead=b[0] + " — ")

    # 6. B2B format
    band(doc, T["s_b2b"], font)
    para(doc, T["b2b_intro"], font)
    for g in T["b2b_groups"]:
        bullet(doc, g[1], font, bold_lead=g[0] + " — ")

    # 7. Speakers
    band(doc, T["s_speakers"], font)
    for sp in T["speakers"]:
        h2(doc, sp[0], font, size=12)
        para(doc, sp[1], font)

    # 8. Contacts
    band(doc, T["s_contact"], font)
    for line in T["contacts"]:
        bullet(doc, line[1], font, bold_lead=line[0] + ": ")

    para(doc, "", font, space_after=10)
    para(doc, T["booklet_foot"], font, size=9, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER)

    path = os.path.join(OUT, filename)
    doc.save(path)
    return path

# ----------------------------------------------------------------------------- CONTENT

exec(open(os.path.join(OUT, "content_ge.py"), encoding="utf-8").read())
exec(open(os.path.join(OUT, "content_ru.py"), encoding="utf-8").read())

if __name__ == "__main__":
    p1 = build_program("Sylfaen", GE, "SPInTechSol_ღონისძიება_პროგრამა_GE.docx")
    p2 = build_booklet("Sylfaen", GE, "SPInTechSol_ღონისძიება_ბუკლეტი_GE.docx")
    p3 = build_program("Calibri", RU, "SPInTechSol_meropriyatie_programma_RU.docx")
    p4 = build_booklet("Calibri", RU, "SPInTechSol_meropriyatie_buklet_RU.docx")
    for p in (p1, p2, p3, p4):
        print("saved:", os.path.basename(p))
